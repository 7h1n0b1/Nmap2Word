#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Author: 7h1n0b1
Description: This program converts the nmap output into a word(docx) file, displayed in table format.
"""

import libnmap
from libnmap.process import NmapProcess
from libnmap.parser import NmapParser
from time import sleep
import tkinter
from tkinter import filedialog
from docx import Document
from docx.shared import Inches

def print_art():
    print(
        '''
        ███    ██ ███    ███  █████  ██████  ███████ ██     ██  ██████  ██████  ██████    
        ████   ██ ████  ████ ██   ██ ██   ██      ██ ██     ██ ██    ██ ██   ██ ██   ██ 
        ██ ██  ██ ██ ████ ██ ███████ ██████   █████  ██  █  ██ ██    ██ ██████  ██   ██ 
        ██  ██ ██ ██  ██  ██ ██   ██ ██      ██      ██ ███ ██ ██    ██ ██   ██ ██   ██ 
        ██   ████ ██      ██ ██   ██ ██      ███████  ███ ███   ██████  ██   ██ ██████  
                                                                                        
                                                                                
        '''
    )

"""
All the scanning related operations 
are defined here
"""

def scan():
    ttype = input("Please select the type of target to scan: \n 1. Manually Enter Target(s) \n 2. Select File Containing Targets (The file should be new line or comma seperated) \n > ")
    if ttype == '1': # When a Single target is to be scanned
        targ = input("Please enter your target to scan: (Ex: mysite.com or 192.168.1.1, scanme.nmap.org or 192.168.1.*) \n > ")
        nswitch = input("\n Please select nmap switches you wish to use: (Ex: -sV -Pn -A) \n > ")
        if targ != '':
            try:
                nmap_proc = NmapProcess(targets=targ, options=nswitch)
                nmap_proc.run_background()
            except PermissionError:
                print("Please make sure Nmap is installed and/or is accessible \n")
                input("Press Enter to exit.")
                exit()
            while nmap_proc.is_running():
                print("Nmap Scan running: ETC: {0} DONE: {1}%".format(nmap_proc.etc,nmap_proc.progress))
                sleep(10) # Adjust this value in order to change how frequently the progress output is printed on terminal.
            print("rc: {0} output: {1}".format(nmap_proc.rc, nmap_proc.summary))
            print(nmap_proc.stderr)
            return nmap_proc.stdout
        else:
            print("Please select correct value \n")
            redo = scan()
    # Triggered when multiple targets are to be slected
    elif ttype == '2':
        floc = input("Input File Path (Press Enter Key to browse your PC): ")
        if floc == '':
            tkwindow = tkinter.Tk()
            tkwindow.withdraw() # Used to Hide the tkinter window.
            filename = filedialog.askopenfilename(initialdir="/", title="Select Your File", filetypes=(("Txt Files", "*.txt"), ("All Files", "*.*")))
            print("The File you selected is: " + filename)
            targets = targ_to_list(filename) 
            nswitch = input("\n Please select nmap switches you wish to use: \n > ")
            try: 
                nmap_proc = NmapProcess(targets=targets, options=nswitch)
                nmap_proc.run_background()
            except PermissionError:
                print("Please make sure Nmap is installed and/or accessible.")
                input("Press Enter to exit.")
                exit()
            while nmap_proc.is_running():
                print("Nmap Scan running: ETC: {0} DONE: {1}%".format(nmap_proc.etc,nmap_proc.progress))
                sleep(10) # Adjust this value in order to change how frequently the progress output is printed on terminal.
            print("rc: {0} output: {1}".format(nmap_proc.rc, nmap_proc.summary))
            print(nmap_proc.stderr)
            return nmap_proc.stdout
        else:
            print(floc)
    else:
        print("Please select a correct option!")
        redo = scan()

""" 
This function is responsible to read XML report, 
extract and return the required output from it.
"""

def outproc(xml_val=''):
    if xml_val != '': # This condition will execute when data is comming from a scan 
        try:
            nmap_report = NmapParser.parse(xml_val)
        except ValueError:
            print(xml_val)
            print("\n The program was not able to process the output properly... \n")
            exit(0)
        except libnmap.parser.NmapParserException:
            print("Something went wrong with the scan and a proper XML report was not generated \n")
            restore = input("We can try to restore the xml output \n Please enter the path where you would like to save the XML file: \n >")
            if restore == '':
                tkwindow = tkinter.Tk()
                tkwindow.withdraw() # Used to Hide the tkinter window.
                while True:
                    try:
                        filename = filedialog.asksaveasfile(initialdir="/", title="Save Your File", filetypes=(("XML Files", "*.xml"), ("All Files", "*.*")))
                        nmap_report = NmapParser.parse_fromfile(filename)
                    except ValueError:
                        print("An error was detected. Was the selected file correct? \n")
                    else:
                        print("XML File Selected: ", restore)
                        break
        tsummary = "Nmap scan summary: {0} \n".format(nmap_report.summary)
        hsummary = "Nmap Scan discovered {0}/{1} hosts up".format(nmap_report.hosts_up, nmap_report.hosts_total)

        final_data = {}
        for _host in nmap_report.hosts: # Loop through all the hosts in XML
            if _host.is_up():   # Check to filter only up hosts
                # print("\n" + "Host: {0} {1}".format(_host.address, " ".join(_host.hostnames)))
                ports = {}
                for s in _host.services: # Loop through all the Ports, Services and State of the current host
                    if s.state == 'open':   # Check to filter only open ports.
                        current_port = {s.port:{"protocol":s.protocol, "service":s.service, "state":s.state}}
                        ports.update(current_port)
                h_data = {_host.address:ports}
                final_data.update(h_data)
        return final_data
    # This will execute when scan was not run and XML file was given directly as an input data
    else: 
        outfile = input("Enter Path to XML File (Press Enter Key To Browse Your PC): ")
        if outfile == '':
            tkwindow = tkinter.Tk()
            tkwindow.withdraw() # Used to Hide the tkinter window.
            while True:
                try:
                    filename = filedialog.askopenfilename(initialdir="/", title="Select Your File", filetypes=(("XML Files", "*.xml"), ("All Files", "*.*")))
                    nmap_report = NmapParser.parse_fromfile(filename)
                except ValueError:
                    print("An error was detected. Was the selected file correct? \n")
                else:
                    print("XML File Selected: ", outfile)
                    break
        else:
            try:
                nmap_report = NmapParser.parse_fromfile(outfile)
                print("XML File Selected: ", outfile)
            except ValueError:
                print("An error was detected. Was the selected file correct? \n")
                redo = outproc()
        tsummary = "Nmap scan summary: {0} \n".format(nmap_report.summary)
        hsummary = "Nmap Scan discovered {0}/{1} hosts up".format(nmap_report.hosts_up, nmap_report.hosts_total)

        final_data = {}
        for _host in nmap_report.hosts: # Loop through all the hosts in XML
            if _host.is_up():   # Check to filter only up hosts
                # print("\n" + "Host: {0} {1}".format(_host.address, " ".join(_host.hostnames)))
                ports = {}
                for s in _host.services: # Loop through all the Ports, Services and State of the current host
                    if s.state == 'open':   # Check to filter only open ports.
                        current_port = {s.port:{"protocol":s.protocol, "service":s.service, "state":s.state}}
                        ports.update(current_port)
                h_data = {_host.address:ports}
                final_data.update(h_data)
        print_scan(nmap_report)
        return final_data

"""
This function is responsible for reading the 
target file, convert it to list and return the value.
"""
def targ_to_list(filename):
    try:
        nmap_file = open(filename, 'r')
    except IOError as err:
        print('File Error: ' + str(err))
    target = []
    while True:
        line = nmap_file.readline() # This will read the file line by line
        if not line:
            break
        if "," in line: # This will run if the file is comma seperated
            line = line.replace(' ', '') # This will remove all the spaces from the file.
            my_line = line.split(",") # Adds the individual values to the "my_line" list
            target.extend(my_line) # The "my_line" list is then added to "target" list
            continue
        target.append(line.strip()) # Strips new line (\n) and add the line to list
    return(target)

"""
Takes a dictionary as input and create the docx output file.
"""
def out_doc(final_input):
    """
    Takes a dictionary as input and looks for nested dictionary
    Find all the keys and print them in Microsoft Word, table format.
    """
    def rec_print(final_input):
        col = 1
        for host, host_info in final_input.items():
            if isinstance(host_info, dict):
                rec_print(host_info)
            else:
                cells = table.cell(row,col)
                cells.text = str(host_info)
                col = col+1

    document = Document()
    document.add_heading('Nmap Output Report', 0)
    document.add_heading('Nman Output', level=1)

    dic_len = len(final_input)
    p = document.add_paragraph("Total {} hosts were found up".format(dic_len))
    for host, host_info in final_input.items():
        p2 = document.add_paragraph()
        para = p2.add_run(str(host))
        para.bold = True
        dic_empty = not bool(host_info) # Returns true if dictionary is empty
        if dic_empty:
            p3 = document.add_paragraph()
            para_itc = p3.add_run("Note: No open ports were found.")
            para_itc.italic = True
            continue # A table will not be created in word, if no open ports were found in the host.
        table = document.add_table(rows=1, cols=4)
        table.style = 'Light Shading Accent 1'
        heading_cells = table.rows[0].cells
        heading_cells[0].text = "Port"
        heading_cells[1].text = "Protocol"
        heading_cells[2].text = "Service"
        heading_cells[3].text = "State"
        row = 0
        for key in host_info:
            cells = table.add_row().cells
            row = row+1
            cells[0].text = str(key)
            rec_print(host_info[key])
        p4 = document.add_paragraph(" ")
    save_dir = input("Enter file name and location to save (Press Enter Key to Browse PC): ")
    if save_dir == '':
        tkwindow = tkinter.Tk()
        tkwindow.withdraw() # Used to Hide the tkinter window.
        while True:
            try:
                save_dir = filedialog.asksaveasfilename(initialdir="/", title="Select folder", filetypes=(("DOCX Files", "*.docx"), ("All Files", "*.*")))
                document.save(str(save_dir))
            except ValueError:
                print("There is something wrong with the path you selected. \n Please select again")
                continue
            except PermissionError:
                print("\n","Oops! It seems the file you are trying to save is either open or you do not have permission to save in that location.", "\n \n")
                input("Press Enter to Try Again...")
                continue
            except FileNotFoundError:
                while True:
                    ans = input("File not found. Do you want to exit without saving your progress? (y/n) > ")
                    if ans == 'y':
                        print("Exiting without saving....")
                        exit()
                    elif ans == 'n':
                        redo = rec_print(final_input)
                        break
                    elif ans != 'y' or 'n':
                        print("Please enter 'y' or 'n'")
                        continue
            else:
                break
        print("File saved at ", save_dir)
    else:
        while True:
            try:
                document.save(save_dir)
            except PermissionError:
                print("\n","Oops! It seems the file you are trying to save is open or you do not have permission to save in that location.", "\n \n", "Please Try Again")
                continue
            except FileNotFoundError:
                while True:
                    ans = input("File not found. Do you want to exit without saving your progress? (y/n) > ")
                    if ans == 'y':
                        print("Exiting without saving....")
                        exit()
                    elif ans == 'n':
                        redo = rec_print(final_input)
                        break
                    elif ans != 'y' or 'n':
                        print("Please enter 'y' or 'n'")
                        continue
            else:
                print("File saved at ", save_dir)
                break
        

"""
Takes the scan output in xml as input and 
gives an output on the terminal.
"""
def print_scan(nmap_report):
    counter = 0
    print(nmap_report)
    try:
        print("Starting Nmap {0} ( http://nmap.org ) at {1}".format(nmap_report.version, nmap_report.started))
    except AttributeError:
        counter = counter + 1
    try:
        for host in nmap_report.hosts:
            if len(host.hostnames):
                tmp_host = host.hostnames.pop()
            else:
                tmp_host = host.address

            print("Nmap scan report for {0} ({1})".format(tmp_host,host.address))
            print("Host is {0}.".format(host.status))
            print("  PORT     STATE         SERVICE")

            for serv in host.services:
                pserv = "{0:>5s}/{1:3s}  {2:12s}  {3}".format(
                        str(serv.port),
                        serv.protocol,
                        serv.state,
                        serv.service)
                if len(serv.banner):
                    pserv += " ({0})".format(serv.banner)
                print(pserv)
    except AttributeError:
        counter = counter+1
    try:
        print(nmap_report.summary)
    except AttributeError:
        counter = counter+1
    if counter > 0:
        print("Note: {} attribute error(s) were detected. The output on the terminal might not be displayed properly".format(counter))

def main():
    print_art()
    taskinp = input("\n Please select your option. \n 1. Scan Your Target (Nmap Should be installed on the host) \n 2. Prepare Report (Requires Nmap's XML output file.) \n > ")
    if taskinp == '1':
       scan_result = scan()
       print("Now parsing the XML output to generate a report.... \n")
       final_output = outproc(scan_result)
       out_doc(final_output)
       parsed = NmapParser.parse(scan_result)
    elif taskinp == '2':
        final_output = outproc()
        out_doc(final_output)
        input("Press Enter to Exit.")
    else:
        print("Please select correct value")
        redo = main()

if __name__ == "__main__":
    main()